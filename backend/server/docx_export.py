"""Markdown to DOCX export matching job_search UI styling."""

import re
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Colors from index.css
OLIVE_GREEN = RGBColor(0x78, 0x8C, 0x5D)  # #788c5d
TEXT_PRIMARY = RGBColor(0x14, 0x14, 0x13)  # #141413
TEXT_MUTED = RGBColor(0x73, 0x72, 0x6C)  # #73726c
LINK_BLUE = RGBColor(0x6A, 0x9B, 0xCC)  # #6a9bcc
FONT_FAMILY = "Calibri"


def _add_bottom_border(paragraph, color_hex: str, width: str = "4"):
    """Add bottom border to paragraph. color_hex without #, e.g. '788C5D'."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), width)  # 1/8 pt units
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _set_margins(doc: Document, inches: float = 0.75) -> None:
    """Set uniform margins on all sections."""
    for section in doc.sections:
        section.top_margin = Inches(inches)
        section.bottom_margin = Inches(inches)
        section.left_margin = Inches(inches)
        section.right_margin = Inches(inches)


def _styled_run(paragraph, text: str, size: Pt, color: RGBColor, bold: bool = False):
    """Add a styled run to paragraph."""
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = size
    run.font.name = FONT_FAMILY
    run.font.color.rgb = color
    return run


def _is_empty_or_hr(line: str) -> bool:
    """Check if line is empty or a horizontal rule."""
    stripped = line.strip()
    return not stripped or bool(re.match(r'^[-*_]{3,}\s*$', stripped))


def markdown_to_docx(markdown: str, output_path: Path) -> None:
    """Convert markdown to DOCX matching job_search CV styling."""
    doc = Document()
    _set_margins(doc)

    # Reset Normal style to tight spacing
    style = doc.styles['Normal']
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(10)
    style.paragraph_format.line_spacing = 1.0

    lines = markdown.split("\n")
    after_h1 = False  # Track if next paragraph is contact line

    for i, line in enumerate(lines):
        # H4: Sub-heading - 10.5pt, bold, muted (check first - most specific)
        if line.startswith("#### "):
            p = doc.add_paragraph()
            _styled_run(p, line[5:].strip(), Pt(10.5), TEXT_MUTED, bold=True)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.keep_with_next = True

        # H3: Job title - 11pt, bold
        elif line.startswith("### "):
            p = doc.add_paragraph()
            _styled_run(p, line[4:].strip(), Pt(11), TEXT_PRIMARY, bold=True)
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.keep_with_next = True

        # H2: Section header - 10.5pt, uppercase, olive green, bordered
        elif line.startswith("## "):
            p = doc.add_paragraph()
            _styled_run(p, line[3:].strip().upper(), Pt(10.5), OLIVE_GREEN, bold=True)
            p.paragraph_format.space_before = Pt(20)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            _add_bottom_border(p, "788C5D", "8")

        # H1: Name - 21pt, bold
        elif line.startswith("# "):
            p = doc.add_paragraph()
            _styled_run(p, line[2:].strip(), Pt(21), TEXT_PRIMARY, bold=True)
            p.space_after = Pt(2)
            after_h1 = True

        # Bullet point
        elif line.strip().startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            _add_formatted_text(p, line.strip()[2:])
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            # Space after last bullet in list
            next_line = lines[i + 1].strip() if i + 1 < len(lines) else ""
            p.paragraph_format.space_after = Pt(0) if next_line.startswith("- ") else Pt(10)

        # Empty line or horizontal rule - reset state
        elif _is_empty_or_hr(line):
            after_h1 = False

        # Regular paragraph (or contact line if after H1)
        else:
            p = doc.add_paragraph()
            _add_formatted_text(p, line, is_contact_line=after_h1)
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            if after_h1:
                p.paragraph_format.space_after = Pt(10)
                _add_bottom_border(p, "E8E6DC", "4")
            else:
                # Add space before next section heading
                next_content = next((ln for ln in lines[i + 1:] if ln.strip()), "")
                if next_content.startswith("#"):
                    p.paragraph_format.space_after = Pt(8)
            after_h1 = False

    doc.save(output_path)


def _add_formatted_text(paragraph, text: str, is_contact_line: bool = False) -> None:
    """Add text with markdown formatting (bold, italic, links)."""
    base_size = Pt(10) if is_contact_line else Pt(10.5)
    base_color = TEXT_MUTED if is_contact_line else TEXT_PRIMARY

    # Named groups: bold_italic, bold, italic, link_text, link_url, plain
    pattern = (
        r"(?:\*\*\*(?P<bold_italic>.+?)\*\*\*)"
        r"|(?:\*\*(?P<bold>.+?)\*\*)"
        r"|(?:\*(?P<italic>.+?)\*)"
        r"|(?:\[(?P<link_text>[^\]]+)\]\((?P<link_url>[^)]+)\))"
        r"|(?P<plain>[^*\[]+)"
    )

    for m in re.finditer(pattern, text):
        if m.group("bold_italic"):
            run = paragraph.add_run(m.group("bold_italic"))
            run.bold, run.italic = True, True
            run.font.color.rgb = TEXT_PRIMARY
        elif m.group("bold"):
            run = paragraph.add_run(m.group("bold"))
            run.bold = True
            run.font.color.rgb = TEXT_PRIMARY
        elif m.group("italic"):
            run = paragraph.add_run(m.group("italic"))
            run.italic = True
            run.font.color.rgb = TEXT_MUTED
            run.font.name = FONT_FAMILY
            run.font.size = Pt(10)  # Italic uses fixed 10pt
            continue
        elif m.group("link_text"):
            run = paragraph.add_run(m.group("link_text"))
            run.font.color.rgb = LINK_BLUE
        elif m.group("plain"):
            run = paragraph.add_run(m.group("plain"))
            run.font.color.rgb = base_color
        else:
            continue

        run.font.name = FONT_FAMILY
        run.font.size = base_size
